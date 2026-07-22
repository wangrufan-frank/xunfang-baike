import json
import tempfile
import unittest
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
CONTENT_PATH = ROOT / "deliverables" / "xunfang-report-speaker-notes.json"


class SpeakerNotesContentTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.payload = json.loads(CONTENT_PATH.read_text(encoding="utf-8"))
        cls.slides = cls.payload["slides"]
        cls.full_text = "\n".join(
            str(value)
            for slide in cls.slides
            for value in slide.values()
            if isinstance(value, str)
        )

    def test_has_exactly_36_ordered_slides(self):
        self.assertEqual(len(self.slides), 36)
        self.assertEqual([item["number"] for item in self.slides], list(range(1, 37)))

    def test_every_slide_has_complete_speaker_fields(self):
        for item in self.slides:
            with self.subTest(slide=item["number"]):
                self.assertTrue(item["title"].strip())
                self.assertGreaterEqual(item["suggested_seconds"], 20)
                self.assertTrue(item["script"].strip())
                self.assertTrue(item["transition"].strip())

    def test_key_facts_and_boundaries_are_preserved(self):
        for term in (
            "6个业务模块",
            "87条内容库存",
            "121个页面",
            "内部学习参考",
            "不替代正式依据",
            "不替代现场履职",
            "当前仅限大队内部试用",
            "未经审查不扩大展示范围",
            "未经批准不扩大展示",
        ):
            with self.subTest(term=term):
                self.assertIn(term, self.full_text)

    def test_live_demo_cues_are_complete(self):
        slide_22 = self.slides[21]
        self.assertIn("切换网站", slide_22["cue"])
        for term in ("首页", "代表性模块", "关键词搜索", "内容详情", "手机端"):
            self.assertIn(term, slide_22["script"] + slide_22["cue"])
        self.assertIn("返回PPT", self.slides[22]["cue"])

    def test_no_placeholders_or_unsupported_outcome_claims(self):
        forbidden = (
            "TBD",
            "TODO",
            "待补充",
            "点击量达到",
            "使用人数达到",
            "满意度达到",
            "全面推广应用",
        )
        for term in forbidden:
            with self.subTest(term=term):
                self.assertNotIn(term, self.full_text)


class SpeakerNotesDocumentTests(unittest.TestCase):
    def test_builder_module_and_document_structure(self):
        try:
            from tools.build_xunfang_report_speaker_notes import build_document
        except ModuleNotFoundError:
            self.skipTest("Word builder is implemented in Task 2")

        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "speaker-notes.docx"
            build_document(CONTENT_PATH, output)
            self.assertTrue(output.exists())
            self.assertGreater(output.stat().st_size, 50_000)

            doc = Document(output)
            text = "\n".join(p.text for p in doc.paragraphs)
            for number in range(1, 37):
                self.assertIn(f"第{number}页", text)
            for term in ("建议时间", "操作提示", "衔接"):
                self.assertIn(term, text)
            self.assertGreater(len(doc.paragraphs), 180)
            header_text = "\n".join(
                paragraph.text
                for section in doc.sections
                for paragraph in section.header.paragraphs
            )
            self.assertIn("巡防百科汇报讲解稿", header_text)


if __name__ == "__main__":
    unittest.main()
