from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


REPORT_NAME = "关于巡防百科网站建设及推广应用情况的报告.docx"


def set_run_font(
    run, latin: str, east_asia: str, size_pt: float, bold: bool = False
):
    """Apply explicit Latin and East Asian font metadata to a run."""
    run.font.name = latin
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), latin)


def configure_page(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(28)


def _configure_style_font(style, latin: str, east_asia: str, size_pt: float, bold=False):
    style.font.name = latin
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), latin)


def configure_styles(document):
    styles = document.styles

    if "Report Title" not in styles:
        report_title = styles.add_style("Report Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        report_title = styles["Report Title"]
    _configure_style_font(
        report_title, "方正小标宋简体", "方正小标宋简体", 22
    )
    report_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_title.paragraph_format.first_line_indent = Pt(0)
    report_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    report_title.paragraph_format.line_spacing = Pt(34)
    report_title.paragraph_format.space_before = Pt(0)
    report_title.paragraph_format.space_after = Pt(24)
    report_title.paragraph_format.keep_with_next = True

    normal = styles["Normal"]
    _configure_style_font(normal, "仿宋_GB2312", "仿宋_GB2312", 16)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(32)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(28)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    heading_1 = styles["Heading 1"]
    _configure_style_font(heading_1, "黑体", "黑体", 16)
    heading_1.font.color.rgb = RGBColor(0, 0, 0)
    heading_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_1.paragraph_format.first_line_indent = Pt(0)
    heading_1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading_1.paragraph_format.line_spacing = Pt(28)
    heading_1.paragraph_format.space_before = Pt(14)
    heading_1.paragraph_format.space_after = Pt(0)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    _configure_style_font(heading_2, "楷体_GB2312", "楷体_GB2312", 16, True)
    heading_2.font.color.rgb = RGBColor(0, 0, 0)
    heading_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_2.paragraph_format.first_line_indent = Pt(0)
    heading_2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading_2.paragraph_format.line_spacing = Pt(28)
    heading_2.paragraph_format.space_before = Pt(8)
    heading_2.paragraph_format.space_after = Pt(0)
    heading_2.paragraph_format.keep_with_next = True

    list_number = styles["List Number"]
    _configure_style_font(list_number, "仿宋_GB2312", "仿宋_GB2312", 16)
    list_number.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    list_number.paragraph_format.left_indent = Mm(12.7)
    list_number.paragraph_format.first_line_indent = Mm(-6.35)
    list_number.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    list_number.paragraph_format.line_spacing = Pt(28)
    list_number.paragraph_format.space_before = Pt(0)
    list_number.paragraph_format.space_after = Pt(0)
    list_number.paragraph_format.widow_control = True

    if "Table Body" not in styles:
        table_body = styles.add_style("Table Body", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_body = styles["Table Body"]
    _configure_style_font(table_body, "仿宋_GB2312", "仿宋_GB2312", 14)
    table_body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table_body.paragraph_format.first_line_indent = Pt(0)
    table_body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    table_body.paragraph_format.line_spacing = Pt(23)
    table_body.paragraph_format.space_before = Pt(0)
    table_body.paragraph_format.space_after = Pt(0)


def add_title(document, text: str):
    paragraph = document.add_paragraph(style="Report Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(34)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(24)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, "方正小标宋简体", "方正小标宋简体", 22)
    return paragraph


def add_body(document, text: str):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(32)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(28)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, "仿宋_GB2312", "仿宋_GB2312", 16)
    return paragraph


def add_heading(document, text: str, level: int = 1):
    if level not in (1, 2):
        raise ValueError("heading level must be 1 or 2")
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", "黑体", 16)
    else:
        set_run_font(run, "楷体_GB2312", "楷体_GB2312", 16, True)
    return paragraph


def add_numbered_points(document, items: list[str]):
    paragraphs = []
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(28)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(item)
        set_run_font(run, "仿宋_GB2312", "仿宋_GB2312", 16)
        paragraphs.append(paragraph)
    return paragraphs


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _format_table_cell(cell, text: str, header: bool = False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.style = "Table Body"
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(23)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, "仿宋_GB2312", "仿宋_GB2312", 14, header)
    if header:
        _set_cell_shading(cell, "F2F4F7")


def add_summary_table(document):
    rows = [
        (
            "内容框架",
            "围绕装备介绍、巡防勤务、警务训练、警情处置、法条规范、走访送教六大业务模块，形成由模块、专题、具体页面构成的三级内容架构。",
            "把分散资料按业务场景归集，便于使用者先按类别定位，再深入查看具体知识。",
        ),
        (
            "查询能力",
            "完成分类浏览、站内搜索及图文、视频辅助展示，建立与内容同步更新的站内索引。",
            "提供按场景进入和按关键词查找两条路径，缩短在多份资料间反复翻找的过程。",
        ),
        (
            "适配验收",
            "完成电脑端和手机端的多终端浏览验收，并对页面跳转、资源引用和内容可读性开展自动检查。",
            "保证主要入口和内容页面在常用终端上均能清晰阅读，为大队内部试用提供基础条件。",
        ),
        (
            "安全治理",
            "开展敏感内容排查和相关内容调整、访问凭据更新、旧凭据失效验证，并对公开法规和公开网页来源进行核验。",
            "为继续完善内容治理和审查把关提供依据，同时明确公开来源核验不替代单位内部业务、保密和发布审批。",
        ),
        (
            "推广反馈",
            "在大队内部围绕日常巡逻和理论学习开展试用，收集可视化、内容完善、图片和视频优化等意见并持续调整。",
            "形成试用、反馈、优化的工作闭环，为后续评估和分层推广积累真实依据。",
        ),
    ]
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    widths_dxa = [1814, 3288, 3742]
    widths_mm = [32, 58, 66]
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "8844")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "B7B7B7")
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for index, (cell, text) in enumerate(
        zip(table.rows[0].cells, ("建设领域", "已完成工作", "实际作用"))
    ):
        cell.width = Mm(widths_mm[index])
        _set_cell_width(cell, widths_dxa[index])
        _format_table_cell(cell, text, header=True)

    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    tr_pr.append(OxmlElement("w:cantSplit"))

    for row_data in rows:
        row = table.add_row()
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for index, (cell, text) in enumerate(zip(row.cells, row_data)):
            cell.width = Mm(widths_mm[index])
            _set_cell_width(cell, widths_dxa[index])
            _format_table_cell(cell, text)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    prefix = paragraph.add_run("— ")
    set_run_font(prefix, "宋体", "宋体", 14)

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")

    page_run = paragraph.add_run()
    set_run_font(page_run, "宋体", "宋体", 14)
    page_run._r.append(field_begin)
    page_run._r.append(instruction)
    page_run._r.append(field_separate)
    display = paragraph.add_run("1")
    set_run_font(display, "宋体", "宋体", 14)
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    display._r.append(field_end)

    suffix = paragraph.add_run(" —")
    set_run_font(suffix, "宋体", "宋体", 14)


def _add_signature(document):
    reporter = document.add_paragraph()
    reporter.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    reporter.paragraph_format.first_line_indent = Pt(0)
    reporter.paragraph_format.left_indent = Mm(82)
    reporter.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    reporter.paragraph_format.line_spacing = Pt(28)
    reporter.paragraph_format.space_before = Pt(28)
    reporter_run = reporter.add_run("（汇报人：________）")
    set_run_font(reporter_run, "仿宋_GB2312", "仿宋_GB2312", 16)

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date.paragraph_format.first_line_indent = Pt(0)
    date.paragraph_format.left_indent = Mm(82)
    date.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    date.paragraph_format.line_spacing = Pt(28)
    date.paragraph_format.space_before = Pt(0)
    date_run = date.add_run("____年__月__日")
    set_run_font(date_run, "仿宋_GB2312", "仿宋_GB2312", 16)


def build_report(output_path: Path) -> None:
    document = Document()
    for section in document.sections:
        configure_page(section)
    configure_styles(document)

    add_title(document, "关于“巡防百科”网站建设及\n推广应用情况的报告")
    add_body(
        document,
        "支队领导：为进一步服务基层巡防实战和民警辅警日常学习，五大队立足工作实际，对分散的巡防业务资料进行系统归集、分类整理和网页化呈现，建设“巡防百科”网站，并在大队内部开展试用。现将有关建设及推广应用情况报告如下。",
    )

    add_heading(document, "一、项目建设背景与初衷")
    add_heading(document, "（一）回应基层知识资料分散的现实问题", 2)
    add_body(
        document,
        "基层巡防工作涉及装备使用、勤务组织、警务训练、警情处置、执法规范和走访送教等多个方面，相关制度规范、培训材料、图文示意和经验做法来源较多、载体不一。过去在需要查询某项业务要点时，往往要在多份文件、课件和个人积累中反复查找，既耗费时间，也不利于形成统一、连续的学习参考。",
    )
    add_heading(document, "（二）推动工作知识由个人积累向集体共享转化", 2)
    add_body(
        document,
        "建设网站的初衷，不是追求形式上的技术展示，而是为大队内部提供一个集中、清晰、便于维护的知识入口，把日常工作中能够沉淀的规范要求、操作要点和培训素材逐步整理为可查询、可复用的工作资源。通过统一内容框架和持续更新机制，减少同类资料重复搜集整理，帮助新警尽快熟悉基本业务，也为骨干民警复习规范、组织训练和开展送教提供参考。",
    )
    add_heading(document, "（三）坚持服务实战、规范稳妥的建设原则", 2)
    add_body(
        document,
        "项目始终把实用性放在首位，围绕基层常见任务设置栏目和查询路径，力求让使用者能够按业务场景进入、按关键词定位、按页面要点学习。同时，坚持安全审慎、边建边用、持续完善，不以未经核实的内容替代正式制度文件，不以网站参考代替现场指挥和依法履职要求。",
    )

    add_heading(document, "二、总体建设思路与服务对象")
    add_heading(document, "（一）形成内容建设闭环", 2)
    add_body(
        document,
        "总体上，项目按照“知识归集、分类整理、网站发布、持续更新”的思路推进。先对已有资料进行梳理，明确主题、适用场景和核心要点；再按照统一结构转化为便于阅读的图文内容；随后通过静态页面生成和站内索引，将内容组织成可浏览、可检索的网站；最后结合试用反馈和制度变化进行复核、补充和优化。",
    )
    add_heading(document, "（二）突出双重服务定位", 2)
    add_body(
        document,
        "一方面，网站定位为新警和转岗人员的入门学习手册，通过模块化、条目化呈现，帮助其建立对巡防业务的基本认识；另一方面，网站定位为一线骨干民警的专业参考入口，便于在日常巡逻、训练备课、警情复盘和走访送教等场景中快速回看相关规范和操作提示。两类需求在内容深度上有所区别，但共同要求入口清楚、表述准确、更新及时。",
    )
    add_heading(document, "（三）明确网站使用边界", 2)
    add_body(
        document,
        "网站现阶段是大队内部的学习参考载体，所列内容应与现行法律法规、上级规范和本单位制度衔接。遇有法律政策调整、上级规定更新或者具体警情处置要求变化时，应以最新正式文件、授权指令和现场指挥为准；对需要业务、保密或者发布审查的内容，必须完成相应程序后方可纳入使用范围。",
    )

    add_heading(document, "三、前期建设工作及阶段性成果")
    add_heading(document, "（一）搭建六大业务模块和三级内容架构", 2)
    add_body(
        document,
        "目前，网站已围绕装备介绍、巡防勤务、警务训练、警情处置、法条规范、走访送教设置六大业务模块，并按照“业务模块、专题分类、具体页面”形成三级内容架构。使用者可先从业务领域进入，再按专题查看具体条目，后续新增内容也能够按同一框架持续归集，避免栏目随意增设、内容相互混杂。",
    )
    add_heading(document, "（二）完善分类浏览和站内查询能力", 2)
    add_body(
        document,
        "网站已形成两条基本查询路径：对目标类别较明确的，可通过首页模块和专题目录逐级浏览；对具体归属不明确的，可输入关键词使用站内搜索定位相关页面。部分内容结合图片和视频进行辅助说明，力求把较长的资料转化为结构清楚、重点可见的学习内容。网站同时支持电脑和手机访问，便于在办公学习和移动阅读等不同场景下使用。",
    )
    add_heading(document, "（三）完成阶段性内容整理和质量检查", 2)
    add_body(
        document,
        "截至2026年7月18日，网站阶段性内容库存为87条，自动检查覆盖121个页面。上述数据用于反映当前建设进度，不代表内容体系已经完备，也不作为推广成效评价。建设过程中已同步开展页面链接、站内索引、资源引用和多终端显示检查，对发现的问题及时调整，确保主要入口和内容页面能够正常浏览。",
    )
    add_body(document, "阶段性建设工作及实际作用概括如下：")
    add_summary_table(document)

    add_heading(document, "四、建设方法与安全管理")
    add_heading(document, "（一）以标准化方式组织内容建设", 2)
    add_body(
        document,
        "在建设方法上，主要开展知识库归集整理、静态页面生成、站内索引、自动检查和多终端验收。知识库归集整理用于统一保存内容来源、主题分类和更新记录；静态页面生成用于把结构化内容转化为稳定、易读的页面；站内索引用于建立关键词与页面之间的对应关系；自动检查和多终端验收用于发现链接失效、资源缺失、页面显示异常等问题。以上方法均服务于内容质量和使用便利，不改变业务审核责任。",
    )
    add_heading(document, "（二）明确现阶段应用范围和持续治理要求", 2)
    add_body(
        document,
        "网站当前预定应用范围为大队内部试用。建设过程中已完成的安全相关工作，主要包括敏感内容排查和相关内容调整、访问凭据更新、旧凭据失效验证。上述工作用于支持现阶段试用管理，但不能替代单位内部审查。对后续新增、修改或者拟扩大使用范围的内容，必须按权限完成业务审查、保密审查和发布审查，未完成相应审查前不作扩展使用。",
    )
    add_heading(document, "（三）严格区分公开来源核验与内部审批", 2)
    add_body(
        document,
        "对涉及法律法规和公开资料的页面，项目已开展公开来源核验，重点确认公开网页的发布主体、正文可访问性以及与页面主题的对应关系。需要特别说明的是，公开来源核验只能证明相关主题或局部内容能够在公开资料中找到参考，不等于网站内容已经完成单位内部发布审批，也不能替代有权人员开展的业务审查、保密审查和发布审查。对仍待审查的内容，应继续保持待审核状态。",
    )
    add_heading(document, "（四）建立问题发现和整改复核机制", 2)
    add_body(
        document,
        "对检查和试用中发现的问题，按照“登记问题、分析原因、修改内容、再次检查”的流程闭环处理。涉及事实表述的，回到制度依据和来源材料进行核对；涉及页面体验的，结合电脑端和手机端实际显示情况优化；涉及安全边界的，优先采取停止展示、删除整改或缩小范围等稳妥措施，整改完成后再进行复核。",
    )

    add_heading(document, "五、大队内部推广应用情况")
    add_heading(document, "（一）推广范围保持在大队内部", 2)
    add_body(
        document,
        "网站目前仅在大队内部推广试用，主要面向民警辅警提供日常工作学习参考，尚未形成跨层级、跨单位推广安排。现阶段工作的重点是验证内容是否贴近基层需要、查询路径是否清晰、手机阅读是否方便，并通过小范围使用尽早发现内容和管理上的不足。",
    )
    add_heading(document, "（二）初步覆盖日常巡逻和理论学习场景", 2)
    add_body(
        document,
        "在日常巡逻准备中，使用者可根据任务需要浏览装备、勤务和常见警情等内容，回顾相关操作提示；在理论学习、训练备课和集中讨论中，可通过法条规范、警务训练和走访送教等模块查找参考材料。网站所发挥的是辅助查询和学习作用，不能代替正式培训、岗位练兵、现场指挥和依法依规处置。",
    )
    add_heading(document, "（三）依据一线反馈持续优化", 2)
    add_body(
        document,
        "试用过程中，大队警力围绕网站可视化、内容完善、图片和视频优化等方面提出了意见建议。项目已根据反馈对栏目呈现、内容结构和多媒体材料进行调整，并继续收集使用中遇到的问题。当前尚未建立完整的量化评估体系，因此对推广情况主要作过程性、定性说明，不以缺乏依据的访问人数、使用频次或者满意度数据评价成效。",
    )

    add_heading(document, "六、当前存在的不足")
    add_heading(document, "（一）内容覆盖仍需持续拓展", 2)
    add_body(
        document,
        "现有内容虽然已经形成基本框架，但与基层巡防任务的广度和更新速度相比仍有差距。部分模块内容深度不够均衡，图片、视频等直观材料仍需补充，一些新装备、新规范和新型任务场景尚未系统纳入，距离完整、成熟的知识体系还有较大提升空间。",
    )
    add_heading(document, "（二）更新维护机制还需制度化", 2)
    add_body(
        document,
        "目前内容更新较多依靠建设人员主动收集和集中整理，职责分工、材料报送、定期复核、失效内容退出等机制还不够完善。随着内容数量增加，如缺少稳定的维护安排，容易出现更新不及时、同类内容表述不一致或者旧规定未及时替换等问题。",
    )
    add_heading(document, "（三）应用效果缺少量化评估", 2)
    add_body(
        document,
        "现阶段已收集到部分使用意见，但尚未围绕查询效率、内容准确性、学习帮助度和问题整改情况建立连续记录，难以对网站在不同岗位、不同场景中的实际作用作出客观比较。后续需要设计简洁、合规、不过度增加基层负担的评价方式，用真实反馈支持优化决策。",
    )
    add_heading(document, "（四）扩大使用范围前仍需严格审查", 2)
    add_body(
        document,
        "网站当前的内容组织和安全措施是按照大队内部试用场景设置的。如后续拟扩大到更高层级或者更多单位使用，内容范围、访问管理、责任分工和安全要求都会发生变化，必须事先开展专项评估，并按权限完成业务、保密和发布审查，不能把内部试用情况直接等同于具备推广条件。",
    )

    add_heading(document, "七、下一步建设计划")
    add_body(
        document,
        "下一阶段，拟坚持需求牵引、内容为本、安全先行、稳步推广，重点做好以下工作。",
    )
    add_numbered_points(
        document,
        [
            "持续丰富业务内容。围绕无人机使用规范、新型装备、常见警情、勤务组织、教育训练等领域补充模块和专题，对已有内容同步开展查漏补缺，优先建设基层使用频率较高、学习需求较集中的条目。",
            "建立月度更新机制。明确资料收集、内容整理、业务复核和发布维护责任，每月对法律政策、上级规范、页面链接和多媒体材料进行一次集中检查，及时标记修订、替换或者停止展示的内容。",
            "推进“每月一学”应用。结合大队教育训练安排，每月选择一个主题，通过网站预习、集中学习、实操训练和复盘讨论相结合的方式使用，推动内容建设与岗位练兵相互促进。",
            "加强审核脱密管理。完善内容分级、来源记录、敏感信息识别和审查留痕，对新增、修改和拟扩大使用范围的内容分别设置审核要求；公开来源核验继续与内部审查相互独立，未经批准不扩大展示范围。",
            "优化移动端阅读体验。根据手机使用中发现的字号、图片、视频、目录和搜索问题，持续调整页面布局和资源大小，使常用信息更加醒目、操作路径更加简洁，同时保留电脑端完整阅读体验。",
            "开展分层试用和反馈评估。在大队内部进一步明确试用岗位和场景，采用问题清单、简短意见收集和定期复盘等方式记录实际效果；待内容、安全和管理条件成熟并履行相应程序后，再研究是否分层扩大试用。",
        ],
    )

    add_heading(document, "八、结语及有关建议")
    add_heading(document, "（一）继续在审核把关前提下开展内部试用", 2)
    add_body(
        document,
        "“巡防百科”网站已完成基础内容框架和查询入口建设，初步具备服务大队内部日常巡逻和理论学习的条件，但仍处于持续完善阶段。建议在严格控制范围、落实内容审核和安全管理的前提下继续开展内部试用，以真实使用需求检验内容价值和操作体验。",
    )
    add_heading(document, "（二）建立内容共建与定期更新机制", 2)
    add_body(
        document,
        "建议结合大队业务分工，探索由相关岗位共同提供资料、业务骨干参与复核、建设人员统一整理维护的协作方式，把网站更新纳入常态学习和训练复盘。对成熟做法及时沉淀，对不适用内容及时退出，逐步形成有人负责、按期复核、问题可追踪的工作机制。",
    )
    add_heading(document, "（三）坚持成熟一项、评估一项、稳妥推进", 2)
    add_body(
        document,
        "后续建设应继续以基层实用为衡量标准，不盲目追求功能和规模。对新增模块、移动端优化和分层试用，均应先明确需求和边界，再开展内容、安全和应用评估。对拟扩大使用范围的事项，严格履行相应程序，确保项目建设始终服务实战、规范有序、安全可控。",
    )
    add_body(document, "以上报告，请审阅。")
    _add_signature(document)

    for section in document.sections:
        configure_page(section)
        footer = section.footer
        paragraph = footer.paragraphs[0]
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
        add_page_number(paragraph)

    document.core_properties.title = "关于“巡防百科”网站建设及推广应用情况的报告"
    document.core_properties.subject = "巡防百科网站项目建设及推广应用情况"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    destination = Path(__file__).resolve().parents[1] / "deliverables" / REPORT_NAME
    build_report(destination)
    print(destination)
