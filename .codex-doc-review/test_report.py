from pathlib import Path
import os
import re
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


REPORT = (
    Path(os.environ["REPORT_PATH"])
    if "REPORT_PATH" in os.environ
    else Path(__file__).resolve().parents[1]
    / "deliverables"
    / "关于巡防百科网站建设及推广应用情况的报告.docx"
)
HEADINGS = [
    "一、项目建设背景与初衷",
    "二、总体建设思路与服务对象",
    "三、前期建设工作及阶段性成果",
    "四、建设方法与安全管理",
    "五、大队内部推广应用情况",
    "六、当前存在的不足",
    "七、下一步建设计划",
    "八、结语及有关建议",
]


class ReportTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.doc = Document(REPORT)
        cls.text = "\n".join(p.text for p in cls.doc.paragraphs)

    def all_document_paragraphs(self):
        for paragraph in self.doc.paragraphs:
            yield paragraph
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
        for section in self.doc.sections:
            for container in (section.header, section.footer):
                yield from container.paragraphs
                for table in container.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from cell.paragraphs

    def text_runs(self, paragraph):
        runs = [run for run in paragraph.runs if run.text.strip()]
        self.assertTrue(runs, f"paragraph has no text runs: {paragraph.text!r}")
        return runs

    def assert_font(self, paragraph, names, size, bold=None):
        for run in self.text_runs(paragraph):
            self.assertIn(run.font.name, names)
            self.assertAlmostEqual(run.font.size.pt, size, places=1)
            if bold is not None:
                self.assertEqual(run.bold, bold)

    def test_required_content(self):
        self.assertEqual(
            self.doc.paragraphs[0].text,
            "关于“巡防百科”网站建设及推广应用情况的报告",
        )
        for heading in HEADINGS:
            self.assertIn(heading, self.text)
        for fact in [
            "六大业务模块",
            "87条",
            "121个页面",
            "大队内部",
            "日常巡逻",
            "理论学习",
            "反馈",
            "优化",
        ]:
            self.assertIn(fact, self.text)

    def test_sensitive_content_absent(self):
        content = "\n".join(paragraph.text for paragraph in self.all_document_paragraphs())
        for forbidden in [
            "口令",
            "密码",
            "SHA-256",
            "SHA-1",
            "SHA-512",
            "MD5",
            "哈希",
            "摘要",
            "Token",
            "token",
            "令牌",
            "Bearer",
            "API Key",
            "API密钥",
            "GitHub",
            "GitLab",
            "仓库",
            "repository",
            "服务器地址",
            "服务器IP",
            "localhost",
            "http://",
            "https://",
            "index.lock",
            "F:\\",
            "C:\\",
            "\\\\",
        ]:
            self.assertNotIn(forbidden, content)
        self.assertIsNone(
            re.search(r"\b(?:\d{1,3}\.){3}\d{1,3}\b", content),
            "report must not disclose a server IP address",
        )

    def test_page_geometry(self):
        for section in self.doc.sections:
            self.assertAlmostEqual(section.page_width.mm, Mm(210).mm, places=1)
            self.assertAlmostEqual(section.page_height.mm, Mm(297).mm, places=1)
            for actual, expected in [
                (section.top_margin.mm, 37),
                (section.bottom_margin.mm, 35),
                (section.left_margin.mm, 28),
                (section.right_margin.mm, 26),
            ]:
                self.assertAlmostEqual(actual, expected, places=1)

    def test_title_and_heading_formatting(self):
        title = self.doc.paragraphs[0]
        self.assertEqual(title.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        for run in self.text_runs(title):
            self.assertEqual(run.font.name, "方正小标宋简体")
            self.assertAlmostEqual(run.font.size.pt, 22, places=1)
            self.assertIsNotNone(run._element.rPr)
            self.assertEqual(
                run._element.rPr.rFonts.get(qn("w:eastAsia")),
                "方正小标宋简体",
            )
        h1s = [
            paragraph for paragraph in self.doc.paragraphs if paragraph.text in HEADINGS
        ]
        self.assertEqual(len(h1s), len(HEADINGS))
        self.assertEqual({paragraph.text for paragraph in h1s}, set(HEADINGS))
        for h1 in h1s:
            for run in self.text_runs(h1):
                self.assertEqual(run.font.name, "黑体")
                self.assertAlmostEqual(run.font.size.pt, 16, places=1)
                self.assertIsNotNone(run._element.rPr)
                self.assertEqual(
                    run._element.rPr.rFonts.get(qn("w:eastAsia")), "黑体"
                )
        h2s = [
            paragraph
            for paragraph in self.doc.paragraphs
            if re.match(r"^（[一二三四五六七八九十]+）", paragraph.text)
        ]
        self.assertTrue(h2s, "report must contain at least one level-2 heading")
        for h2 in h2s:
            self.assert_font(h2, ["楷体_GB2312"], 16, bold=True)

    def test_body_paragraph_formatting(self):
        style = self.doc.styles["Normal"]
        self.assertAlmostEqual(style.font.size.pt, Pt(16).pt, places=1)
        self.assertEqual(style.font.name, "仿宋_GB2312")
        self.assertAlmostEqual(style.paragraph_format.line_spacing.pt, Pt(28).pt, places=1)
        bodies = [
            paragraph
            for paragraph in self.doc.paragraphs
            if paragraph.text
            and paragraph.style.name == "Normal"
            and paragraph != self.doc.paragraphs[0]
            and paragraph.text not in HEADINGS
            and not re.match(r"^（[一二三四五六七八九十]+）", paragraph.text)
            and not paragraph.text.startswith("（汇报人：")
            and not paragraph.text.startswith("____年")
        ]
        self.assertTrue(bodies, "report must contain at least one body paragraph")
        for body in bodies:
            self.assertEqual(body.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assert_font(body, ["仿宋_GB2312"], 16)
            self.assertEqual(body.paragraph_format.line_spacing_rule, WD_LINE_SPACING.EXACTLY)
            self.assertAlmostEqual(body.paragraph_format.line_spacing.pt, Pt(28).pt, places=1)
            self.assertAlmostEqual(body.paragraph_format.first_line_indent.pt, Pt(32).pt, places=1)

    def test_footer_page_number_formatting(self):
        for section in self.doc.sections:
            footer = section.footer.paragraphs[0]
            self.assertEqual(footer.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertTrue(footer.text.startswith("— "))
            self.assertTrue(footer.text.endswith(" —"))
            self.assertIn("PAGE", footer._p.xml)

    def test_summary_table_geometry_and_row_height_contract(self):
        self.assertEqual(len(self.doc.tables), 1)
        table = self.doc.tables[0]
        expected_widths = [1814, 3288, 3742]
        self.assertEqual(len(table.columns), len(expected_widths))

        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        self.assertIsNotNone(tbl_w)
        self.assertEqual(tbl_w.get(qn("w:type")), "dxa")
        self.assertEqual(int(tbl_w.get(qn("w:w"))), sum(expected_widths))
        tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
        self.assertIsNotNone(tbl_ind)
        self.assertEqual(tbl_ind.get(qn("w:type")), "dxa")
        self.assertEqual(tbl_ind.get(qn("w:w")), "120")

        grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        self.assertEqual(grid_widths, expected_widths)
        self.assertEqual(sum(grid_widths), int(tbl_w.get(qn("w:w"))))

        header_pr = table.rows[0]._tr.get_or_add_trPr()
        header = header_pr.first_child_found_in("w:tblHeader")
        self.assertIsNotNone(header)
        self.assertEqual(header.get(qn("w:val")), "true")
        for row in table.rows:
            tr_pr = row._tr.trPr
            self.assertTrue(
                tr_pr is None or tr_pr.first_child_found_in("w:trHeight") is None,
                "summary table must not use fixed row heights",
            )
            self.assertEqual(len(row.cells), len(expected_widths))
            for cell, expected_width in zip(row.cells, expected_widths):
                tc_w = cell._tc.tcPr.first_child_found_in("w:tcW")
                self.assertIsNotNone(tc_w)
                self.assertEqual(tc_w.get(qn("w:type")), "dxa")
                self.assertEqual(int(tc_w.get(qn("w:w"))), expected_width)

    def test_signature_placeholders(self):
        self.assertIn("（汇报人：________）", self.text)
        self.assertIn("____年__月__日", self.text)


if __name__ == "__main__":
    unittest.main()
