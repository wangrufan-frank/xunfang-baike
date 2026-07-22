import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


NAVY = RGBColor(0x1B, 0x31, 0x68)
RED = RGBColor(0xC7, 0x20, 0x2F)
INK = RGBColor(0x22, 0x29, 0x33)
GRAY = RGBColor(0x66, 0x72, 0x80)
LIGHT_BLUE = "EEF3F8"
LIGHT_RED = "FBECEE"
LIGHT_GRAY = "F3F5F7"
BODY_FONT = "思源宋体 CN Medium"
HEAD_FONT = "思源黑体 CN Heavy"


def set_run_font(run, name, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph, color, size=18, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    borders.append(left)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])
    set_run_font(run, HEAD_FONT, 9, GRAY)


def configure_page(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)


def apply_document_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.4
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if "Document Title" not in styles:
        styles.add_style("Document Title", WD_STYLE_TYPE.PARAGRAPH)
    title = styles["Document Title"]
    title.font.name = HEAD_FONT
    title._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), HEAD_FONT)
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = NAVY
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)

    if "Section Title" not in styles:
        styles.add_style("Section Title", WD_STYLE_TYPE.PARAGRAPH)
    section_title = styles["Section Title"]
    section_title.font.name = HEAD_FONT
    section_title._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), HEAD_FONT)
    section_title.font.size = Pt(18)
    section_title.font.bold = True
    section_title.font.color.rgb = NAVY
    section_title.paragraph_format.space_before = Pt(8)
    section_title.paragraph_format.space_after = Pt(10)
    section_title.paragraph_format.keep_with_next = True

    if "Slide Title" not in styles:
        styles.add_style("Slide Title", WD_STYLE_TYPE.PARAGRAPH)
    slide_title = styles["Slide Title"]
    slide_title.font.name = HEAD_FONT
    slide_title._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), HEAD_FONT)
    slide_title.font.size = Pt(15)
    slide_title.font.bold = True
    slide_title.font.color.rgb = NAVY
    slide_title.paragraph_format.space_before = Pt(14)
    slide_title.paragraph_format.space_after = Pt(6)
    slide_title.paragraph_format.keep_with_next = True
    slide_title.paragraph_format.keep_together = True

    if "Speaker Body" not in styles:
        styles.add_style("Speaker Body", WD_STYLE_TYPE.PARAGRAPH)
    speaker = styles["Speaker Body"]
    speaker.font.name = BODY_FONT
    speaker._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    speaker.font.size = Pt(12)
    speaker.font.color.rgb = INK
    speaker.paragraph_format.first_line_indent = Cm(0.85)
    speaker.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    speaker.paragraph_format.line_spacing = 1.4
    speaker.paragraph_format.space_after = Pt(6)
    speaker.paragraph_format.widow_control = True


def set_running_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("巡防百科汇报讲解稿")
    set_run_font(run, HEAD_FONT, 9, NAVY, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    add_page_number(p)


def add_cover(doc, root):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("汇 报 讲 解 稿")
    set_run_font(run, HEAD_FONT, 12, RED, bold=True)

    p = doc.add_paragraph(style="Document Title")
    p.add_run("关于“巡防百科”网站建设及\n推广应用情况的汇报")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("36页PPT配套完整逐字稿")
    set_run_font(run, HEAD_FONT, 13, GRAY)

    cover_image = root / "img" / "auth-patrol-night.jpg"
    if cover_image.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        p.add_run().add_picture(str(cover_image), width=Cm(15.7))

    for label, value in (
        ("汇报对象", "大队领导、支队领导"),
        ("汇报人", "王汝凡"),
        ("时间", "2026年7月"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{label}：{value}")
        set_run_font(run, BODY_FONT, 12, INK)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_usage_guide(doc, payload):
    doc.add_paragraph("使用说明", style="Section Title")

    p = doc.add_paragraph(style="Speaker Body")
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run(
        "本讲解稿与36页PPT逐页对应，完整讲解预计约35—40分钟。现场可根据时间删减正文，但建议保留关键数据、现场演示步骤和安全边界。"
    )

    for label, value, color in (
        ("页面对应", "每节标题中的页码与PPT页码一致。", NAVY),
        ("操作提示", "红色提示用于翻页、指图、切换网站和返回PPT。", RED),
        ("衔接", "每页末尾提供自然过渡语，可直接用于翻页。", GRAY),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.space_after = Pt(5)
        label_run = p.add_run(f"{label}｜")
        set_run_font(label_run, HEAD_FONT, 11, color, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, BODY_FONT, 11, INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"建议总时长：{payload['estimated_minutes']}　｜　汇报对象：{payload['audience']}")
    set_run_font(run, HEAD_FONT, 10, GRAY)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def format_duration(seconds):
    minutes, remain = divmod(seconds, 60)
    if minutes and remain:
        return f"约{minutes}分{remain}秒"
    if minutes:
        return f"约{minutes}分钟"
    return f"约{remain}秒"


def add_info_paragraph(doc, label, text, fill, accent):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.right_indent = Cm(0.15)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.25
    set_paragraph_shading(p, fill)
    set_paragraph_left_border(p, accent)
    label_run = p.add_run(f"{label}｜")
    set_run_font(label_run, HEAD_FONT, 10.5, RGBColor.from_string(accent), bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, BODY_FONT, 10.5, INK)
    return p


def add_slide_section(doc, item):
    section_starts = {1, 8, 12, 18, 23, 28, 31, 36}
    if item["number"] in section_starts and item["number"] != 1:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    heading = doc.add_paragraph(style="Slide Title")
    heading.add_run(f"第{item['number']}页｜{item['title']}")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(f"建议时间｜{format_duration(item['suggested_seconds'])}")
    set_run_font(run, HEAD_FONT, 9.5, GRAY, bold=True)

    cue = item.get("cue", "").strip()
    if cue:
        add_info_paragraph(doc, "操作提示", cue, LIGHT_RED, "C7202F")

    for block in item["script"].split("\n\n"):
        p = doc.add_paragraph(style="Speaker Body")
        p.add_run(block.strip())

    add_info_paragraph(doc, "衔接", item["transition"], LIGHT_BLUE, "1B3168")


def build_document(content_path, output_path):
    content_path = Path(content_path)
    output_path = Path(output_path)
    payload = json.loads(content_path.read_text(encoding="utf-8"))
    if len(payload["slides"]) != 36:
        raise ValueError("Speaker notes must contain exactly 36 slides")

    doc = Document()
    configure_page(doc.sections[0])
    apply_document_styles(doc)
    set_running_header_footer(doc.sections[0])

    root = Path(__file__).resolve().parents[1]
    add_cover(doc, root)
    add_usage_guide(doc, payload)
    for item in payload["slides"]:
        add_slide_section(doc, item)

    core = doc.core_properties
    core.title = payload["title"]
    core.subject = "36页PPT配套完整逐字讲解稿"
    core.author = "王汝凡"
    core.keywords = "巡防百科, 汇报, 讲解稿, PPT"
    core.comments = "根据确认的36页汇报PPT制作"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Build the Xunfang report speaker notes DOCX")
    parser.add_argument("--content", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build_document(args.content, args.output)
    print(f"DOCX_OK {args.output}")


if __name__ == "__main__":
    main()
