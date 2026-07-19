"""从 Word 文档自动提取法规全文，写入 legal-documents.json。

用法：
  1. 将 .docx 文件放到 data/legal-texts/ 目录下
  2. 文件名应与法规 id 匹配，例如：
     - 治安管理处罚法.docx → zhian-guanli-chufa-fa
     - 人民警察法.docx     → renmin-jingcha-fa
     - 居民身份证法.docx   → jumin-shenfenzheng-fa
     或者文件名包含关键词即可（脚本会自动匹配）
  3. 运行:  python tools/ingest_law_docx.py
          python tools/ingest_law_docx.py --check   # 只预览不写入
"""

from argparse import ArgumentParser
from collections import OrderedDict
from pathlib import Path
import json
import re
import sys

try:
    from docx import Document as DocxDocument
except ImportError:
    print('请先安装 python-docx: pip install python-docx', file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parents[1]
TEXT_DIR = ROOT / 'data' / 'legal-texts'
LEDGER_PATH = ROOT / 'data' / 'legal-documents.json'

# 文件名关键词 → law id
KEYWORD_MAP = OrderedDict([
    ('治安管理处罚', 'zhian-guanli-chufa-fa'),
    ('人民警察法', 'renmin-jingcha-fa'),
    ('居民身份证法', 'jumin-shenfenzheng-fa'),
    ('行政案件程序', 'xingzheng-anji-chengxu-guiding'),
    ('警械和武器条例', 'jingxie-wuqi-tiaoli'),
    ('现场制止', 'xianchang-zhizhi-guicheng'),
    ('操作规程', 'xianchang-zhizhi-guicheng'),  # fallback
])


def find_docx_files() -> dict[str, Path]:
    """扫描 TEXT_DIR 下的 .docx 文件，匹配到 law id。"""
    matches = {}
    for path in sorted(TEXT_DIR.glob('*.docx')):
        name = path.stem
        for keyword, law_id in KEYWORD_MAP.items():
            if keyword in name:
                matches[law_id] = path
                print(f'📎 {path.name} → {law_id}')
                break
        else:
            print(f'⚠️  无法识别: {path.name}，请将文件名改为包含法规关键词')
    return matches


def extract_paragraphs(doc) -> list[str]:
    """从 Word 文档提取所有段落文本（按顺序，跳过空行）。"""
    texts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
    # 也检查表格中的文本（有些文档把法条放在表格里）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    texts.append(text)
    return texts


def parse_law_paragraphs(paragraphs: list[str]) -> list[dict]:
    """从段落文本列表解析为 chapters[] 结构。"""
    CHAPTER_RE = re.compile(
        r'^第[一二三四五六七八九十百千]+章\s*[，、．.\s]*(.+)'
    )
    ARTICLE_RE = re.compile(
        r'^第([一二三四五六七八九十百千零\d]+)条\s*[，、．.\s]*(.*)'
    )

    chapters = []
    current_chapter = None
    current_article = None
    preamble = []  # 法规标题、题注等前置内容

    for line in paragraphs:
        line = line.strip()
        if not line:
            continue

        chapter_match = CHAPTER_RE.match(line)
        article_match = ARTICLE_RE.match(line)

        if chapter_match:
            # 保存前一条
            if current_article is not None and current_chapter is not None:
                current_chapter['articles'].append(current_article)
                current_article = None
            if current_chapter is not None:
                # 跳过目录中的章标题（没有条文的就是目录条目）
                if current_chapter['articles']:
                    chapters.append(current_chapter)
                # else: TOC entry, discard

            title = chapter_match.group(1).strip() if chapter_match.group(1) else ''
            current_chapter = {
                'number': line.split()[0] if line.split() else line[:4],
                'title': title or line.split()[0] if line.split() else '',
                'articles': [],
            }
            continue

        if article_match:
            if current_chapter is None:
                current_chapter = {
                    'number': '',
                    'title': '',
                    'articles': [],
                }

            if current_article is not None:
                current_chapter['articles'].append(current_article)

            num_str = article_match.group(1)
            try:
                number = _cn_to_int(num_str)
            except ValueError:
                number = len(current_chapter['articles']) + 1

            current_article = {
                'number': number,
                'label': f'第{num_str}条',
                'paragraphs': [],
            }

            # 条标题后的文字如果有，作为第一款
            rest = article_match.group(2).strip()
            if rest:
                current_article['paragraphs'].append(rest)
            continue

        # 普通段落 → 属于当前条文的正文
        if current_article is not None:
            current_article['paragraphs'].append(line)
        elif current_chapter is None and not chapters:
            # 在第一章之前的内容（标题、题注等）→ 跳过
            preamble.append(line)

    # 收尾
    if current_article is not None and current_chapter is not None:
        current_chapter['articles'].append(current_article)
    if current_chapter is not None:
        chapters.append(current_chapter)

    return chapters


def _cn_to_int(s: str) -> int:
    """中文数字 → int。如 '一百一十九' → 119。"""
    if s.isdigit():
        return int(s)
    cn = {'零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
          '五': 5, '六': 6, '七': 7, '八': 8, '九': 9,
          '十': 10, '百': 100, '千': 1000}
    result = 0
    temp = 0
    for ch in s:
        if ch == '百':
            temp *= 100
            result += temp
            temp = 0
        elif ch == '十':
            temp = (temp or 1) * 10
            result += temp
            temp = 0
        else:
            val = cn.get(ch)
            if val is None:
                raise ValueError(f'无法解析数字: {s!r} 中的 {ch!r}')
            temp += val
    return result + temp


def process_one(law_id: str, docx_path: Path, ledger: dict, *, check_only: bool = False) -> bool:
    """处理一个 docx 文件。"""
    print(f'\n📖 处理: {docx_path.name}')

    try:
        doc = DocxDocument(str(docx_path))
    except Exception as e:
        print(f'❌ 无法打开 {docx_path}: {e}')
        return False

    paragraphs = extract_paragraphs(doc)
    print(f'   提取到 {len(paragraphs)} 个段落')

    chapters = parse_law_paragraphs(paragraphs)
    total_articles = sum(len(ch['articles']) for ch in chapters)

    if total_articles == 0:
        print(f'❌ 未解析到任何条文。请确认 Word 文档中包含"第X条"格式的条文。')
        return False

    print(f'   解析结果: {len(chapters)} 章, {total_articles} 条')
    for ch in chapters[:5]:
        arts = [a['label'] for a in ch['articles'][:3]]
        print(f'   {ch["number"]} {ch["title"]}: {", ".join(arts)}{" ..." if len(ch["articles"]) > 3 else ""}')
    if len(chapters) > 5:
        print(f'   ... 共 {len(chapters)} 章')

    if check_only:
        return True

    # 更新 ledger
    doc_entry = next((d for d in ledger['documents'] if d['id'] == law_id), None)
    if doc_entry is None:
        print(f'❌ legal-documents.json 中找不到 id={law_id}')
        return False

    doc_entry['chapters'] = chapters
    doc_entry['partial'] = False
    if not doc_entry.get('verified_at'):
        doc_entry['verified_at'] = '2026-07-19'

    # 同时更新 .txt 文件（方便后续人工校对）
    txt_path = TEXT_DIR / f'{law_id}.txt'
    txt_content = _render_to_txt(doc_entry)
    txt_path.write_text(txt_content, encoding='utf-8')
    print(f'   📝 同步更新: {txt_path.name}')

    return True


def _render_to_txt(doc: dict) -> str:
    """将法规数据渲染为可读的 .txt 格式。"""
    lines = [
        f'# {doc["title"]}',
        f'# 官方来源: {doc.get("source_url", "")}',
        f'# 自动提取于: 2026-07-19（请人工校对后删除本注释）',
        '',
    ]
    for ch in doc['chapters']:
        lines.append(f'## {ch["number"]} {ch["title"]}'.strip())
        lines.append('')
        for art in ch['articles']:
            lines.append(f'### {art["label"]}')
            for para in art['paragraphs']:
                lines.append(para)
            lines.append('')
    return '\n'.join(lines)


def main():
    parser = ArgumentParser(description=__doc__)
    parser.add_argument('--check', action='store_true', help='仅预览，不写入')
    args = parser.parse_args()

    # 扫描 docx 文件
    matches = find_docx_files()
    if not matches:
        print('未找到 .docx 文件。请将法规 Word 文档放到:')
        print(f'  {TEXT_DIR}')
        print('文件名示例: 治安管理处罚法.docx, 人民警察法.docx')
        return 1

    # 加载 ledger
    try:
        ledger = json.loads(LEDGER_PATH.read_text(encoding='utf-8'))
    except (OSError, json.JSONDecodeError) as e:
        print(f'❌ 无法加载 {LEDGER_PATH}: {e}')
        return 1

    # 逐个处理
    all_ok = True
    processed = 0
    for law_id, docx_path in matches.items():
        ok = process_one(law_id, docx_path, ledger, check_only=args.check)
        if ok:
            processed += 1
        else:
            all_ok = False

    if processed == 0:
        print('\n⚠️  没有成功处理的文件。')
        return 1

    if not args.check and processed > 0:
        LEDGER_PATH.write_text(
            json.dumps(ledger, ensure_ascii=False, indent=2) + '\n',
            encoding='utf-8',
        )
        print(f'\n📦 已更新 {LEDGER_PATH} ({processed} 部法规)')
        print(f'👉 下一步: python tools/build_legal_pages.py')
        print(f'👉 然后验证: python -m unittest discover -s tests -p "test_*.py" -v')

    return 0 if all_ok else 1


if __name__ == '__main__':
    sys.exit(main())
